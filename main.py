import os
import io
import json
import re
import tempfile
from datetime import datetime
from pathlib import Path
import logging
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from functools import wraps

# Flask imports
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify
from flask_sqlalchemy import SQLAlchemy
import sqlite3

# Image processing imports
import pytesseract
from PIL import Image
import cv2
import numpy as np
from paddleocr import PaddleOCR
from textblob import TextBlob

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import fitz  # PyMuPDF

# LangChain imports - Gemini 
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.schema import Document as LangChainDocument

from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv("SECRET_KEY")
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///answer_evaluator.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['KNOWLEDGE_BASE_FOLDER'] = 'knowledge_base'
app.config['VECTOR_STORE_FOLDER'] = 'vector_stores'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['KNOWLEDGE_BASE_FOLDER'], exist_ok=True)
os.makedirs(app.config['VECTOR_STORE_FOLDER'], exist_ok=True)

# Initialize database
db = SQLAlchemy(app)

# API Configuration
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")


# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    role = db.Column(db.String(20), nullable=False)  # 'teacher' or 'student'
    class_name = db.Column(db.String(50))  # For students
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Subject(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    teacher_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    class_name = db.Column(db.String(50), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Book(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    subject_id = db.Column(db.Integer, db.ForeignKey('subject.id'), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)

class QuestionPaper(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    subject_id = db.Column(db.Integer, db.ForeignKey('subject.id'), nullable=False)
    subject = db.relationship('Subject', backref='question_papers')
    class_name = db.Column(db.String(50), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    total_questions = db.Column(db.Integer, default=0)
    questions_json = db.Column(db.Text)  # Store questions as JSON
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)
    

class AnswerSheet(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    question_paper_id = db.Column(db.Integer, db.ForeignKey('question_paper.id'), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    total_score = db.Column(db.Float, default=0.0)
    evaluation_json = db.Column(db.Text)  # Store detailed evaluation
    submitted_at = db.Column(db.DateTime, default=datetime.utcnow)
    evaluated = db.Column(db.Boolean, default=False)

# OCR and Document Processing Classes
class OCRProcessor:
    def __init__(self, engine: str = "tesseract"):
        self.engine = engine.lower()
        if self.engine == "paddle":
            self.ocr = PaddleOCR(use_angle_cls=True, lang="en")
        elif self.engine == "gemini":
            if not GOOGLE_API_KEY:
                raise ValueError("GOOGLE_API_KEY required for Gemini OCR")

    def preprocess_image(self, image: np.ndarray) -> np.ndarray:
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        thresh = cv2.adaptiveThreshold(
            blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        kernel = np.ones((1, 1), np.uint8)
        processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        return processed

    def extract_text_from_image(self, image_bytes: bytes) -> str:
        try:
            if self.engine == "tesseract":
                image = Image.open(io.BytesIO(image_bytes))
                opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                processed = self.preprocess_image(opencv_image)
                custom_config = r"--oem 3 --psm 6"
                text = pytesseract.image_to_string(processed, config=custom_config)

            elif self.engine == "paddle":
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                    tmp_file.write(image_bytes)
                    tmp_path = tmp_file.name

                results = self.ocr.ocr(tmp_path, cls=True)
                texts = []
                if results and len(results) > 0:
                    for line in results[0]:
                        if len(line) > 1:
                            text = line[1][0]
                            confidence = line[1][1]
                            if confidence > 0.3:
                                texts.append(text)
                text = " ".join(texts)
                os.remove(tmp_path)

            elif self.engine == "gemini":
                # Gemini Vision API for OCR
                import google.generativeai as genai
                genai.configure(api_key=GOOGLE_API_KEY)
                model = genai.GenerativeModel('gemini-2.0-flash')
                
                image = Image.open(io.BytesIO(image_bytes))
                response = model.generate_content([
                    "Extract all handwritten text from this image. Return only the extracted text without any additional commentary.",
                    image
                ])
                text = response.text


            return self._clean_text(text)
        except Exception as e:
            logger.error(f"OCR error ({self.engine}): {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        return " ".join(lines)

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_scanned_pdf(self, file_path: str, ocr_processor) -> str:
        full_text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                # High resolution improves handwriting OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(5, 5))

                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                tmp_path = tmp.name
                tmp.close()  # Close before saving to avoid "Permission denied"

                # ✅ Save pixmap after closing the handle
                pix.save(tmp_path)

                # ✅ Read bytes for OCR
                with open(tmp_path, "rb") as f:
                    img_bytes = f.read()

                # ✅ Pass valid image bytes to the selected OCR engine (Gemini or Paddle)
                page_text = ocr_processor.extract_text_from_image(img_bytes)
                if page_text:
                    full_text += page_text + "\n"

                os.remove(tmp.name)
            doc.close()
            return full_text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from scanned PDF: {e}", exc_info=True)
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""

class KnowledgeBaseManager:
    def __init__(self, subject_id: int):
        self.subject_id = subject_id
        self.embeddings = self._initialize_embeddings()
        self.vector_store = None
        self.documents = []
        self.vector_store_path = os.path.join(
            app.config['VECTOR_STORE_FOLDER'], 
            f"subject_{subject_id}"
        )

    def _initialize_embeddings(self):
        if GOOGLE_API_KEY:
            try:
                embeddings = GoogleGenerativeAIEmbeddings(
                    model="models/embedding-001",
                    google_api_key=GOOGLE_API_KEY
                )
                return embeddings
            except Exception as e:
                logger.error(f"Failed to initialize embeddings: {e}")
                raise
        else:
            raise RuntimeError("GOOGLE_API_KEY not found")

    def load_documents(self, file_paths: list) -> None:
        doc_processor = DocumentProcessor()
        
        for file_path in file_paths:
            file_path = Path(file_path)
            
            if not file_path.exists():
                continue
            
            if file_path.suffix.lower() == '.pdf':
                text = doc_processor.extract_text_from_pdf(str(file_path))
                if not text.strip() or len(text.strip()) < 50:
                    ocr = OCRProcessor("tesseract")
                    text = doc_processor.extract_text_from_scanned_pdf(str(file_path), ocr)
            elif file_path.suffix.lower() == '.docx':
                text = doc_processor.extract_text_from_docx(str(file_path))
            else:
                continue
            
            if text:
                chunks = doc_processor.text_splitter.split_text(text)
                for i, chunk in enumerate(chunks):
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={"source": str(file_path), "chunk": i}
                    )
                    self.documents.append(doc)

    def create_vector_store(self) -> None:
        if not self.documents:
            raise ValueError("No documents loaded")
        
        self.vector_store = FAISS.from_documents(self.documents, self.embeddings)
        os.makedirs(self.vector_store_path, exist_ok=True)
        self.vector_store.save_local(self.vector_store_path)

    def load_vector_store(self) -> None:
        if os.path.exists(self.vector_store_path):
            self.vector_store = FAISS.load_local(
                self.vector_store_path, 
                self.embeddings,
                allow_dangerous_deserialization=True
            )

class AnswerEvaluator:
    def __init__(self, subject_id: int):
        self.subject_id = subject_id
        self.kb_manager = KnowledgeBaseManager(subject_id)
        
        if not GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY required for evaluation")
        
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash", 
            google_api_key=GOOGLE_API_KEY,
            temperature=0.3
        )
        
        self.evaluation_prompt = PromptTemplate(
            input_variables=["context", "question", "student_answer"],
            template="""
            You are an expert teacher evaluating student answers.

            Context from textbooks:
            {context}

            Question: {question}
            Student's Answer: {student_answer}

            Evaluate based on:
            1. Correctness (0-20 points)
            2. Completeness (0-20 points)
            3. Understanding (0-20 points)
            4. Terminology (0-20 points)
            5. Clarity (0-20 points)

            Provide evaluation in JSON format:
            {{
                "score": <total out of 100>,
                "scores": {{
                    "correctness": <score>,
                    "completeness": <score>,
                    "understanding": <score>,
                    "terminology": <score>,
                    "clarity": <score>
                }},
                "feedback": "<detailed feedback>",
                "strengths": ["<strength 1>", "<strength 2>"],
                "improvements": ["<improvement 1>", "<improvement 2>"],
                "model_answer": "<brief model answer>"
            }}
            """
        )

    def evaluate_answer(self, question: str, student_answer: str) -> dict:
        try:
            if self.kb_manager.vector_store is None:
                try:
                    self.kb_manager.load_vector_store()
                except Exception as load_error:
                    logger.warning(f"Could not load vector store: {load_error}")
                    # Continue without context
                    context = "No reference material available for this subject."
        
            # Get context from knowledge base if available
            if self.kb_manager.vector_store is not None:
            
                retriever = self.kb_manager.vector_store.as_retriever(search_kwargs={"k": 3})
                relevant_docs = retriever.invoke(question)
                context = "\n\n".join([doc.page_content for doc in relevant_docs])
            
                if not context.strip():
                    context = "No reference material found."
            else:
                context = "No reference material available for this subject."
            
            formatted_prompt = self.evaluation_prompt.format(
                context=context,
                question=question,
                student_answer=student_answer
            )
            
            response = self.llm.invoke(formatted_prompt)
            
            try:
                response_text = response.content
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)

                
                if json_match:
                    evaluation = json.loads(json_match.group())
                    return evaluation
                else:
                    return {
                        "score": 0,
                        "scores": {
                            "correctness": 0,
                            "completeness": 0,
                            "understanding": 0,
                            "terminology": 0,
                            "clarity": 0
                                },
                        "feedback": "Evaluation failed to generate proper structured response.",
                        "strengths": [],
                        "improvements": [],
                        "model_answer": "Reference answer unavailable."
                            }

            except json.JSONDecodeError:
                return {
                        "score": 0,
                        "scores": {
                            "correctness": 0,
                            "completeness": 0,
                            "understanding": 0,
                            "terminology": 0,
                            "clarity": 0
                                },
                        "feedback": "Evaluation failed to generate proper structured response.",
                        "strengths": [],
                        "improvements": [],
                        "model_answer": "Reference answer unavailable."
                            }
        except Exception as e:
            logger.error(f"Evaluation error: {e}")
            return {
                        "score": 0,
                        "scores": {
                            "correctness": 0,
                            "completeness": 0,
                            "understanding": 0,
                            "terminology": 0,
                            "clarity": 0
                                },
                        "feedback": "Evaluation failed to generate proper structured response.",
                        "strengths": [],
                        "improvements": [],
                        "model_answer": "Reference answer unavailable."
                            }

# Helper functions
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def teacher_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        user = User.query.get(session['user_id'])
        if user.role != 'teacher':
            flash('Access denied. Teachers only.', 'error')
            return redirect(url_for('student_dashboard'))
        return f(*args, **kwargs)
    return decorated_function

def student_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        user = User.query.get(session['user_id'])
        if user.role != 'student':
            flash('Access denied. Students only.', 'error')
            return redirect(url_for('teacher_dashboard'))
        return f(*args, **kwargs)
    return decorated_function

# Routes
@app.route('/')
def index():
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        if user.role == 'teacher':
            return redirect(url_for('teacher_dashboard'))
        else:
            return redirect(url_for('student_dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            
            if user.role == 'teacher':
                return redirect(url_for('teacher_dashboard'))
            else:
                return redirect(url_for('student_dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        role = request.form.get('role')
        class_name = request.form.get('class_name') if role == 'student' else None
        
        if User.query.filter_by(username=username).first():
            flash('Username already exists', 'error')
            return redirect(url_for('register'))
        
        hashed_password = generate_password_hash(password)
        new_user = User(
            username=username,
            email=email,
            password=hashed_password,
            role=role,
            class_name=class_name
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')

# Add these routes to the main Flask app

@app.route('/teacher/dashboard')
@teacher_required
def teacher_dashboard():
    user = User.query.get(session['user_id'])
    subjects = Subject.query.filter_by(teacher_id=user.id).all()
    
    # Get statistics
    stats = {}
    for subject in subjects:
        question_papers = QuestionPaper.query.filter_by(
            subject_id=subject.id,
            class_name=subject.class_name
        ).count()
        
        answer_sheets = db.session.query(AnswerSheet).join(QuestionPaper).filter(
            QuestionPaper.subject_id == subject.id,
            QuestionPaper.class_name == subject.class_name
        ).count()
        
        stats[subject.id] = {
            'question_papers': question_papers,
            'answer_sheets': answer_sheets
        }
    
    return render_template('teacher/dashboard.html', subjects=subjects, stats=stats)

@app.route('/teacher/subject/create', methods=['GET', 'POST'])
@teacher_required
def create_subject():
    if request.method == 'POST':
        name = request.form.get('name')
        class_name = request.form.get('class_name')
        
        subject = Subject(
            name=name,
            teacher_id=session['user_id'],
            class_name=class_name
        )
        
        db.session.add(subject)
        db.session.commit()
        
        flash(f'Subject "{name}" created successfully for class {class_name}!', 'success')
        return redirect(url_for('teacher_dashboard'))
    
    return render_template('teacher/create_subject.html')

@app.route('/teacher/subject/<int:subject_id>/books', methods=['GET', 'POST'])
@teacher_required
def manage_books(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    
    if subject.teacher_id != session['user_id']:
        flash('Access denied', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    if request.method == 'POST':
        files = request.files.getlist('books')
        
        if not files:
            flash('No files uploaded', 'error')
            return redirect(url_for('manage_books', subject_id=subject_id))
        
        uploaded_books = []
        file_paths = []
        
        for file in files:
            if file and file.filename:
                filename = secure_filename(file.filename)
                subject_folder = os.path.join(
                    app.config['KNOWLEDGE_BASE_FOLDER'],
                    f"subject_{subject_id}"
                )
                os.makedirs(subject_folder, exist_ok=True)
                
                file_path = os.path.join(subject_folder, filename)
                file.save(file_path)
                
                book = Book(
                    filename=filename,
                    subject_id=subject_id,
                    file_path=file_path
                )
                db.session.add(book)
                uploaded_books.append(filename)
                file_paths.append(file_path)
        
        db.session.commit()
        
        # Create knowledge base
        try:
            kb_manager = KnowledgeBaseManager(subject_id)
            kb_manager.load_documents(file_paths)
            kb_manager.create_vector_store()
            
            flash(f'Successfully uploaded {len(uploaded_books)} books and created knowledge base!', 'success')
        except Exception as e:
            flash(f'Books uploaded but knowledge base creation failed: {str(e)}', 'warning')
        
        return redirect(url_for('manage_books', subject_id=subject_id))
    
    books = Book.query.filter_by(subject_id=subject_id).all()
    return render_template('teacher/manage_books.html', subject=subject, books=books)

@app.route('/teacher/subject/<int:subject_id>/question-paper/upload', methods=['GET', 'POST'])
@teacher_required
def upload_question_paper(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    
    if subject.teacher_id != session['user_id']:
        flash('Access denied', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    if request.method == 'POST':
        title = request.form.get('title')
        file = request.files.get('question_paper')
        
        if not file or not file.filename:
            flash('No file uploaded', 'error')
            return redirect(url_for('upload_question_paper', subject_id=subject_id))
        
        filename = secure_filename(file.filename)
        qp_folder = os.path.join(
            app.config['UPLOAD_FOLDER'],
            'question_papers',
            f"subject_{subject_id}"
        )
        os.makedirs(qp_folder, exist_ok=True)
        
        file_path = os.path.join(qp_folder, filename)
        file.save(file_path)
        
        # Extract questions
        doc_processor = DocumentProcessor()
        
        if filename.endswith('.pdf'):
            text = doc_processor.extract_text_from_pdf(file_path)
        elif filename.endswith('.docx'):
            text = doc_processor.extract_text_from_docx(file_path)
        else:
            flash('Unsupported file format', 'error')
            return redirect(url_for('upload_question_paper', subject_id=subject_id))
        
        # Parse questions
        questions = re.split(r"(?:Ques|Question|Q)\.?\s*\d+[:.)]?", text, flags=re.IGNORECASE)
        question_list = []
        for i, q in enumerate(questions):
            q = q.strip()
            if q:
                question_list.append({"number": i, "text": q})
        
        question_paper = QuestionPaper(
            title=title,
            subject_id=subject_id,
            class_name=subject.class_name,
            file_path=file_path,
            total_questions=len(question_list),
            questions_json=json.dumps(question_list)
        )
        
        db.session.add(question_paper)
        db.session.commit()
        
        flash(f'Question paper uploaded successfully! Found {len(question_list)} questions.', 'success')
        return redirect(url_for('view_question_papers', subject_id=subject_id))
    
    return render_template('teacher/upload_question_paper.html', subject=subject)

@app.route('/teacher/subject/<int:subject_id>/question-papers')
@teacher_required
def view_question_papers(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    
    if subject.teacher_id != session['user_id']:
        flash('Access denied', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    question_papers = QuestionPaper.query.filter_by(
        subject_id=subject_id,
        class_name=subject.class_name
    ).all()
    
    return render_template('teacher/question_papers.html', subject=subject, question_papers=question_papers)

@app.route('/teacher/subject/<int:subject_id>/students')
@teacher_required
def view_students(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    
    if subject.teacher_id != session['user_id']:
        flash('Access denied', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    # Get all students in this class
    students = User.query.filter_by(role='student', class_name=subject.class_name).all()
    
    # Get answer sheets for each student
    student_data = []
    total_scores = []
    
    for student in students:
        answer_sheets = db.session.query(AnswerSheet).join(QuestionPaper).filter(
            AnswerSheet.student_id == student.id,
            QuestionPaper.subject_id == subject_id,
            QuestionPaper.class_name == subject.class_name
        ).all()
        
        total_score = sum([sheet.total_score for sheet in answer_sheets if sheet.evaluated])
        evaluated_count = len([sheet for sheet in answer_sheets if sheet.evaluated])
        avg_score = total_score / evaluated_count if evaluated_count > 0 else 0
        
        student_data.append({
            'student': student,
            'answer_sheets': answer_sheets,
            'avg_score': round(avg_score, 2),
            'total_submissions': len(answer_sheets),
            'evaluated_submissions': evaluated_count
        })
        
        if evaluated_count > 0:
            total_scores.append(avg_score)
    
    # Calculate class statistics
    class_average = round(sum(total_scores) / len(total_scores), 2) if total_scores else 0
    highest_score = round(max(total_scores), 2) if total_scores else 0
    
    return render_template('teacher/students.html', 
                         subject=subject, 
                         student_data=student_data,
                         class_average=class_average,
                         highest_score=highest_score)

@app.route('/teacher/evaluate/<int:answer_sheet_id>', methods=['GET', 'POST'])
@teacher_required
def evaluate_answer_sheet(answer_sheet_id):
    answer_sheet = AnswerSheet.query.get_or_404(answer_sheet_id)
    question_paper = QuestionPaper.query.get(answer_sheet.question_paper_id)
    subject = Subject.query.get(question_paper.subject_id)
    
    if subject.teacher_id != session['user_id']:
        flash('Access denied', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    if request.method == 'POST':
        books = Book.query.filter_by(subject_id=subject.id).count()
        if books == 0:
            flash('Please upload reference books for this subject before evaluation.', 'warning')
            return redirect(url_for('manage_books', subject_id=subject.id))
        
        # CHECK IF VECTOR STORE EXISTS
        vector_store_path = os.path.join(
            app.config['VECTOR_STORE_FOLDER'], 
            f"subject_{subject.id}"
        )
        if not os.path.exists(vector_store_path):
            flash('Knowledge base not found. Please re-upload books to create it.', 'warning')
            return redirect(url_for('manage_books', subject_id=subject.id))
        
        ocr_mode = request.form.get('ocr_mode', 'paddle')
        
        # Extract text from answer sheet
        if ocr_mode == 'gemini':
            ocr_processor = OCRProcessor('gemini')
        else:
            ocr_processor = OCRProcessor('paddle')
        
        doc_processor = DocumentProcessor()
        
        if answer_sheet.file_path.endswith('.pdf'):
            extracted_text = doc_processor.extract_text_from_scanned_pdf(
                answer_sheet.file_path, 
                ocr_processor
            )
        else:
            with open(answer_sheet.file_path, 'rb') as f:
                image_bytes = f.read()
            extracted_text = ocr_processor.extract_text_from_image(image_bytes)
        
        # Parse answers
        answers = [
            ans.strip()
            for ans in re.split(r"(?:A\s*n\s*s\w*\s*\d+[:.)]?)", extracted_text, flags=re.IGNORECASE)
            if ans.strip()
        ]
        
        # Get questions
        questions = json.loads(question_paper.questions_json)

        
        # Evaluate each answer
        evaluator = AnswerEvaluator(subject.id)
        evaluations = []
        total_score = 0
        
        for i, question in enumerate(questions):
            answer = answers[i] if i < len(answers) else ""
            evaluation = evaluator.evaluate_answer(question["text"], answer)
            
            evaluation.update({
                "question_number": question["number"],
                "question": question["text"],
                "student_answer": answer
            })
            evaluations.append(evaluation)
            
            if "score" in evaluation and not evaluation.get("error"):
                total_score += evaluation["score"]
        
        # Calculate average
        avg_score = round(total_score / len(questions), 2) if questions else 0
        
        # Save evaluation
        answer_sheet.evaluation_json = json.dumps(evaluations)
        answer_sheet.total_score = avg_score
        answer_sheet.evaluated = True
        db.session.commit()
        
        flash('Answer sheet evaluated successfully!', 'success')
        return redirect(url_for('view_evaluation', answer_sheet_id=answer_sheet_id))
    
    student = User.query.get(answer_sheet.student_id)
    
    return render_template('teacher/evaluate.html', 
                         answer_sheet=answer_sheet,
                         question_paper=question_paper,
                         subject=subject,
                         student=student)

@app.route('/teacher/evaluation/<int:answer_sheet_id>')
@teacher_required
def view_evaluation(answer_sheet_id):
    answer_sheet = AnswerSheet.query.get_or_404(answer_sheet_id)
    question_paper = QuestionPaper.query.get(answer_sheet.question_paper_id)
    subject = Subject.query.get(question_paper.subject_id)
    student = User.query.get(answer_sheet.student_id)
    
    if subject.teacher_id != session['user_id']:
        flash('Access denied', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    evaluations = json.loads(answer_sheet.evaluation_json) if answer_sheet.evaluation_json else []
    
    return render_template('teacher/view_evaluation.html',
                         answer_sheet=answer_sheet,
                         question_paper=question_paper,
                         subject=subject,
                         student=student,
                         evaluations=evaluations)

# Add these routes to the main Flask app

@app.route('/student/dashboard')
@student_required
def student_dashboard():
    user = User.query.get(session['user_id'])
    
    # Get available question papers for student's class
    question_papers = QuestionPaper.query.filter_by(class_name=user.class_name).all()
    
    # Get student's submissions
    answer_sheets = AnswerSheet.query.filter_by(student_id=user.id).all()
    
    # Calculate statistics
    evaluated_sheets = [sheet for sheet in answer_sheets if sheet.evaluated]
    total_score = sum([sheet.total_score for sheet in evaluated_sheets])
    avg_score = round(total_score / len(evaluated_sheets), 2) if evaluated_sheets else 0
    
    # Get rank in class
    all_students = User.query.filter_by(role='student', class_name=user.class_name).all()
    student_scores = []
    
    for student in all_students:
        student_sheets = AnswerSheet.query.filter_by(student_id=student.id).all()
        evaluated = [sheet for sheet in student_sheets if sheet.evaluated]
        if evaluated:
            student_avg = sum([sheet.total_score for sheet in evaluated]) / len(evaluated)
            student_scores.append({
                'student_id': student.id,
                'avg_score': student_avg
            })
    
    # Sort by score descending
    student_scores.sort(key=lambda x: x['avg_score'], reverse=True)
    
    rank = None
    for i, score_data in enumerate(student_scores):
        if score_data['student_id'] == user.id:
            rank = i + 1
            break
    
    return render_template('student/dashboard.html',
                         question_papers=question_papers,
                         answer_sheets=answer_sheets,
                         avg_score=avg_score,
                         rank=rank,
                         total_students=len(student_scores))

@app.route('/student/question-paper/<int:qp_id>')
@student_required
def view_question_paper(qp_id):
    user = User.query.get(session['user_id'])
    question_paper = QuestionPaper.query.get_or_404(qp_id)
    
    if question_paper.class_name != user.class_name:
        flash('Access denied', 'error')
        return redirect(url_for('student_dashboard'))
    
    questions = json.loads(question_paper.questions_json) if question_paper.questions_json else []
    
    return render_template('student/view_question_paper.html',
                         question_paper=question_paper,
                         questions=questions)

@app.route('/student/submit-answer/<int:qp_id>', methods=['GET', 'POST'])
@student_required
def submit_answer(qp_id):
    user = User.query.get(session['user_id'])
    question_paper = QuestionPaper.query.get_or_404(qp_id)
    
    if question_paper.class_name != user.class_name:
        flash('Access denied', 'error')
        return redirect(url_for('student_dashboard'))
    
    if request.method == 'POST':
        file = request.files.get('answer_sheet')
        
        if not file or not file.filename:
            flash('No file uploaded', 'error')
            return redirect(url_for('submit_answer', qp_id=qp_id))
        
        filename = secure_filename(file.filename)
        answer_folder = os.path.join(
            app.config['UPLOAD_FOLDER'],
            'answer_sheets',
            f"student_{user.id}"
        )
        os.makedirs(answer_folder, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_path = os.path.join(answer_folder, f"{timestamp}_{filename}")
        file.save(file_path)
        
        answer_sheet = AnswerSheet(
            student_id=user.id,
            question_paper_id=qp_id,
            file_path=file_path
        )
        
        db.session.add(answer_sheet)
        db.session.commit()
        
        flash('Answer sheet submitted successfully!', 'success')
        return redirect(url_for('student_dashboard'))
    
    return render_template('student/submit_answer.html', question_paper=question_paper)

@app.route('/student/my-submissions')
@student_required
def my_submissions():
    user = User.query.get(session['user_id'])
    answer_sheets = AnswerSheet.query.filter_by(student_id=user.id).all()
    
    submissions_data = []
    for sheet in answer_sheets:
        qp = QuestionPaper.query.get(sheet.question_paper_id)
        subject = Subject.query.get(qp.subject_id)
        
        submissions_data.append({
            'answer_sheet': sheet,
            'question_paper': qp,
            'subject': subject
        })
    
    return render_template('student/my_submissions.html', submissions=submissions_data)

@app.route('/student/result/<int:answer_sheet_id>')
@student_required
def view_result(answer_sheet_id):
    user = User.query.get(session['user_id'])
    answer_sheet = AnswerSheet.query.get_or_404(answer_sheet_id)
    
    if answer_sheet.student_id != user.id:
        flash('Access denied', 'error')
        return redirect(url_for('student_dashboard'))
    
    if not answer_sheet.evaluated:
        flash('This answer sheet has not been evaluated yet.', 'warning')
        return redirect(url_for('my_submissions'))
    
    question_paper = QuestionPaper.query.get(answer_sheet.question_paper_id)
    subject = Subject.query.get(question_paper.subject_id)
    evaluations = json.loads(answer_sheet.evaluation_json) if answer_sheet.evaluation_json else []
    
    return render_template('student/view_result.html',
                         answer_sheet=answer_sheet,
                         question_paper=question_paper,
                         subject=subject,
                         evaluations=evaluations)

@app.route('/student/rankings')
@student_required
def view_rankings():
    user = User.query.get(session['user_id'])
    
    # Get all students in the class
    all_students = User.query.filter_by(role='student', class_name=user.class_name).all()
    
    student_rankings = []
    
    for student in all_students:
        answer_sheets = AnswerSheet.query.filter_by(student_id=student.id).all()
        evaluated_sheets = [sheet for sheet in answer_sheets if sheet.evaluated]
        
        if evaluated_sheets:
            total_score = sum([sheet.total_score for sheet in evaluated_sheets])
            avg_score = round(total_score / len(evaluated_sheets), 2)
            
            student_rankings.append({
                'student': student,
                'avg_score': avg_score,
                'total_submissions': len(evaluated_sheets)
            })
    
    # Sort by average score descending
    student_rankings.sort(key=lambda x: x['avg_score'], reverse=True)
    
    # Add rank
    for i, ranking in enumerate(student_rankings):
        ranking['rank'] = i + 1
    
    return render_template('student/rankings.html', 
                         rankings=student_rankings,
                         current_user_id=user.id)

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, port=5000)